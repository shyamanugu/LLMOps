#!/usr/bin/env python3
"""Generate the Afni Enterprise LLMOps written proposal (Word .docx).

Requires: python-docx  (pip install python-docx)
Output:   proposal/Afni-LLMOps-Proposal.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x12, 0x1F, 0x3D)
INDIGO= RGBColor(0x1B, 0x3A, 0x6B)
TEAL  = RGBColor(0x00, 0x7A, 0x7A)
AMBER = RGBColor(0xB5, 0x74, 0x00)
GRAY  = RGBColor(0x50, 0x5A, 0x6A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT  = "Segoe UI"

doc = Document()

# base styles
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor(0x22, 0x28, 0x33)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, size, col in [(1, 17, NAVY), (2, 13.5, INDIGO), (3, 11.5, TEAL)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = FONT
    st.font.size = Pt(size)
    st.font.color.rgb = col
    st.font.bold = True
    st.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
    st.paragraph_format.space_after = Pt(4)


def _shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def para(text="", size=10.5, bold=False, italic=False, color=None, align=None,
         space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.name = FONT
        if color:
            r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.name = FONT
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(10.5)


def h1(t): return doc.add_heading(t, level=1)
def h2(t): return doc.add_heading(t, level=2)
def h3(t): return doc.add_heading(t, level=3)


def make_table(headers, rows, widths=None, header_fill="1B3A6B"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        _shade(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(htext)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        r.font.name = FONT
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                _shade(cells[ci], "EEF1F7")
            p = cells[ci].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            r.font.name = FONT
            if ci == 0:
                r.bold = True
                r.font.color.rgb = NAVY
    if widths:
        for ci, wdt in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = wdt
    para("", space_after=4)
    return t


def callout(title, text):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shade(cell, "EAF3F3")
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    r.bold = True; r.font.color.rgb = TEAL; r.font.name = FONT; r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(10.5)
    para("", space_after=4)

print("Building document...")

# ---- page setup + footer with page number ----------------------------------
sec = doc.sections[0]
sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

def add_page_number(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run("Evoke Technologies  ·  Enterprise LLMOps for Afni  ·  Confidential  ·  Page ")
fr.font.size = Pt(8); fr.font.color.rgb = GRAY; fr.font.name = FONT
add_page_number(footer_p)

# ============================================================ COVER PAGE
para("", space_after=60)
para("ENTERPRISE PROPOSAL  ·  2026", 12, bold=True, color=TEAL,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para("Enterprise LLMOps Platform for Afni", 30, bold=True, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para("Industrializing multi-agent Generative AI — Voice AI for contact centers and "
     "AI-driven HR recruitment on one governed platform.", 13, italic=True, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
para("Prepared by Evoke Technologies · GenAI Architecture Practice", 12, bold=True, color=INDIGO,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Author: Shyam — Senior GenAI Architect, embedded at Afni", 11, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("Draft v1.0  ·  Confidential — for Afni & Evoke review", 10, color=GRAY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

make_table(["Field", "Detail"],
           [["Client", "Afni, Inc. — global BPO & customer engagement"],
            ["Prepared by", "Evoke Technologies (Shyam, Senior GenAI Architect)"],
            ["Document", "Enterprise LLMOps Proposal"],
            ["Version / Status", "v1.0 / Draft for review"],
            ["Classification", "Confidential"],
            ["Date", "2026"]],
           widths=[Inches(1.8), Inches(4.6)])
doc.add_page_break()

# ============================================================ CONTENTS
h1("Contents")
toc = [
    "1. Executive Summary", "2. Afni Business Context & Opportunity",
    "3. Platform Architecture", "4. Multi-Agent Systems Design",
    "5. Use Case 1 — Voice AI for Contact Centers",
    "6. Use Case 2 — AI-Driven HR Recruitment",
    "7. LLMOps Lifecycle & Toolchain", "8. Responsible AI & Governance",
    "9. Security, Privacy & Compliance", "10. Observability & FinOps",
    "11. Operating Model & Team", "12. Implementation Roadmap",
    "13. Business Case & ROI", "14. Risks & Mitigations", "15. Recommendation & Next Steps",
]
for item in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(item); r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = INDIGO
    if item.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
        r.bold = True
doc.add_page_break()

# ============================================================ 1. EXEC SUMMARY
h1("1. Executive Summary")
para("Afni operates in a business where Generative AI is no longer optional. As a global "
     "business process outsourcing and customer-engagement provider, Afni's economics are "
     "driven by two things this proposal targets directly: the cost and quality of high-volume "
     "voice interactions, and the cost and speed of hiring the people who handle them. "
     "AI-native competitors are resetting client expectations on price, availability and "
     "consistency. The opportunity is to move first — as an AI-enabled partner, not a repriced one.")
para("This proposal recommends that Afni stand up an enterprise-grade LLMOps platform on "
     "Microsoft Azure AI Foundry and use it to deliver two flagship, multi-agent use cases:")
bullet("customer and agent-facing Voice AI across the contact-center estate; and", "Voice AI — ")
bullet("AI-driven HR recruitment for Afni's own high-volume, multi-country hiring.", "HR recruitment — ")
para("The strategy is deliberately platform-first. Rather than building disconnected pilots, "
     "Afni builds the foundations — orchestration, retrieval, evaluation, guardrails, "
     "observability, security and governance — once, so that every subsequent use case "
     "(subrogation automation, quality analytics, a knowledge assistant) reuses them and ships "
     "in weeks rather than quarters. Both flagship use cases run on the same multi-agent "
     "pattern: an orchestrator that routes work to specialist agents for intent, knowledge/RAG, "
     "actions, compliance, sentiment, escalation and summarization.")
callout("The core idea:",
        "Give Afni one secure, governed platform to build, evaluate, deploy, govern and "
        "continuously improve fleets of cooperating AI agents — with the same operational rigor "
        "Afni already applies to running contact centers.")
para("Delivery follows a Crawl → Walk → Run roadmap over roughly nine to twelve months, "
     "beginning with a four-week Phase 0 that establishes the Azure landing zone, security "
     "baseline and governance, followed by pilots that prove value early. Responsible AI, "
     "human-in-the-loop control and compliance (PCI-DSS, HIPAA, TCPA, SOC 2, GDPR, and "
     "hiring-fairness law such as EEOC and NYC Local Law 144) are designed in from day one. "
     "All financial figures in this document are illustrative placeholders to be replaced with "
     "Afni's actuals during discovery.")
para("Recommendation: approve Phase 0 to lock the foundations and launch the agent-assist and "
     "HR-screening pilots.", bold=True, color=NAVY)

# ============================================================ 2. BUSINESS CONTEXT
h1("2. Afni Business Context & Opportunity")
para("Afni, Inc. is a global BPO and customer-engagement provider founded in 1936, "
     "headquartered in Bloomington, Illinois, with a workforce of roughly 3,400+ professionals "
     "operating across the United States, Mexico and the Philippines, complemented by the "
     "Afni@Home remote program. (Public-source figures; to be confirmed with Afni.)")
h2("2.1 Service lines")
make_table(["Service line", "What it involves", "GenAI relevance"],
           [["Acquisition & Growth", "Sales and customer acquisition", "Outbound assist, lead qualification"],
            ["Care & Retention", "Customer service and loyalty", "Voice agents, agent-assist, self-service"],
            ["Collections", "Receivables and recovery", "Compliant reminders, negotiation assist"],
            ["P&C Insurance", "Insurance support incl. subrogation", "Document AI, claims & subrogation automation"],
            ["Gainshare model", "Outcome-based partnerships", "Efficiency directly improves shared margin"]],
           widths=[Inches(1.7), Inches(2.5), Inches(2.4)])
h2("2.2 Strategic pressures on BPOs")
bullet("margin and labor-cost pressure as AI-native providers reset price expectations;", "Economics — ")
bullet("persistent agent attrition and ramp cost keep recruiting and training in permanent motion;", "People — ")
bullet("clients now expect 24/7, instant, consistent and compliant interactions;", "Experience — ")
bullet("first movers become strategic AI partners; laggards get repriced.", "Disruption — ")
h2("2.3 GenAI opportunity map")
make_table(["Candidate use case", "Business impact", "Effort", "Priority"],
           [["Agent-assist copilot (voice)", "High", "Medium", "Now (Phase 1)"],
            ["HR recruitment automation", "High", "Low–Med", "Now (Phase 1)"],
            ["Autonomous voice agent", "High", "Medium", "Phase 2"],
            ["Post-call QA & analytics", "Medium–High", "Low", "Phase 2"],
            ["Subrogation / document automation", "High", "Medium", "Phase 3"],
            ["Knowledge assistant (internal)", "Medium", "Low", "Phase 3"]],
           widths=[Inches(2.4), Inches(1.5), Inches(1.1), Inches(1.4)])

# ============================================================ 3. ARCHITECTURE
h1("3. Platform Architecture")
para("The platform is a layered reference architecture on Azure AI Foundry. Each layer has a "
     "clear responsibility and named Azure services, so that use cases compose from shared "
     "building blocks rather than bespoke stacks.")
make_table(["Layer", "Responsibility", "Key Azure services"],
           [["Experience & channels", "Where users interact", "CCaaS + Azure Communication Services, web/chat, agent desktop, ATS/HR portals"],
            ["Orchestration & agents", "Coordinate multi-agent work", "Azure AI Agent Service, Semantic Kernel / AutoGen (Microsoft Agent Framework)"],
            ["Models & AI services", "Reasoning, speech, safety", "Azure OpenAI GPT-4o / mini, gpt-realtime, AI Speech, Content Safety"],
            ["Knowledge & RAG", "Grounded answers", "Azure AI Search (hybrid + semantic), AI Document Intelligence"],
            ["Data & integration", "State, records, connectors", "Cosmos DB, Data Lake / Fabric, Azure SQL, API Management, Functions"],
            ["DevOps & LLMOps", "Build, test, ship", "Prompt flow & registry, AI evaluation SDK, GitHub Actions / Azure DevOps, Container Apps / AKS"],
            ["Security & governance", "Trust & control", "Entra ID, Key Vault, Defender for Cloud, Purview, private endpoints"],
            ["Observability & FinOps", "See & cost everything", "Azure Monitor, App Insights, OpenTelemetry, token metering"]],
           widths=[Inches(1.7), Inches(1.9), Inches(3.0)])
h2("3.1 Environments & isolation")
para("Separate dev, test and production environments follow the Azure Cloud Adoption Framework "
     "landing-zone pattern. Traffic stays on private networking (VNet + private endpoints) with "
     "no public egress for sensitive data; identities are managed through Entra ID with "
     "least-privilege RBAC and managed identities; secrets live in Key Vault.")
h2("3.2 Model serving & AI gateway")
para("Azure API Management sits in front of model endpoints as an AI gateway, providing token "
     "metering, per-consumer quotas and throttling, semantic caching, and routing/fallback "
     "across models. This is also the enforcement point for cost control and abuse protection.")
h2("3.3 Alternatives considered")
para("AWS Bedrock Agents and Google Vertex AI Agent Builder are credible alternatives. Azure is "
     "recommended as primary given Afni's likely Microsoft enterprise footprint and the "
     "maturity of Azure AI Foundry for regulated, multi-agent workloads with integrated agent, "
     "safety and governance tooling. The gateway abstraction keeps the platform model-portable.")

# ============================================================ 4. MULTI-AGENT
h1("4. Multi-Agent Systems Design")
para("Both flagship use cases are built on one pattern: a supervisor/orchestrator agent that "
     "interprets a request and routes work to specialist agents, each with a narrow, testable "
     "responsibility. Probabilistic agents are wrapped in deterministic guardrails so that "
     "policy — disclosures, PII handling, do-not-say/must-say — is enforced, not merely hoped for.")
make_table(["Specialist agent", "Responsibility"],
           [["Intent / Router", "Classify the request and route to the right specialist"],
            ["Knowledge / RAG", "Retrieve grounded, cited answers from policy and knowledge bases"],
            ["Action / Tooling", "Call systems of record (CRM, HRIS/ATS, billing) via secure tools"],
            ["Compliance / Guardrail", "Enforce disclosures, PII redaction, TCPA/EEOC and fairness rules"],
            ["Sentiment / Emotion", "Detect frustration or escalation signals"],
            ["Escalation / Handoff", "Warm-transfer to a human with full context"],
            ["Summarization / QA", "Produce summaries, dispositions and QA scores"]],
           widths=[Inches(2.0), Inches(4.4)])
h2("4.1 Orchestration patterns")
para("The platform supports supervisor-orchestrator, sequential, concurrent, hand-off, "
     "group-chat, and reflection/critic patterns, with human-in-the-loop checkpoints for "
     "consequential decisions. Agents use tool/function calling to act; short-term memory holds "
     "conversation context while long-term memory (Cosmos DB) persists profiles and history.")
h2("4.2 Frameworks")
para("Orchestration is implemented with Semantic Kernel and/or AutoGen (converging in the "
     "Microsoft Agent Framework) and hosted via Azure AI Agent Service. The same seven "
     "specialists serve both use cases — in voice they answer callers; in HR they move "
     "candidates through the funnel — proving the platform's reusability.")

# ============================================================ 5. VOICE
h1("5. Use Case 1 — Voice AI for Contact Centers")
para("Voice is Afni's highest-volume, highest-cost channel. The use case is delivered in three "
     "modes, sequenced from lowest to highest risk.")
h2("5.1 Three modes")
bullet("live transcription, next-best-action, knowledge surfacing, sentiment and compliance "
       "nudges, and automatic summary/disposition for human reps. Lowest risk, fastest value.",
       "Agent-assist copilot — ")
bullet("natural, sub-second speech-to-speech handling of containable call types (FAQs, "
       "verification, appointments, payment reminders), with warm handoff to a human when "
       "needed.", "Autonomous voice agent — ")
bullet("every call summarized and scored for quality and compliance, replacing sampled QA with "
       "100% coverage and generating coaching insights.", "Post-call analytics & QA — ")
h2("5.2 Call flow")
para("Caller → telephony/CCaaS → realtime speech-to-text → orchestrator → specialist agents "
     "(RAG + tools) → systems of record → spoken response (TTS) or human handoff. Content "
     "Safety, PII redaction and TCPA/PCI guardrails run on every turn, within a sub-second "
     "latency budget. Payment moments trigger PCI pause/mask.", italic=False)
h2("5.3 KPIs")
make_table(["KPI", "Baseline (today)", "Target with Voice AI *"],
           [["Call containment / deflection", "Manual IVR", "20–40% of eligible call types"],
            ["Average Handle Time", "Program baseline", "15–25% reduction"],
            ["QA coverage", "2–10% sampled", "100% automated"],
            ["Compliance adherence", "Sampled review", "Monitored on every call"],
            ["Agent ramp time", "Weeks", "Materially shorter with copilot"]],
           widths=[Inches(2.4), Inches(1.9), Inches(2.1)])
para("* Illustrative ranges — replaced with Afni actuals during discovery.", 9, italic=True, color=GRAY)
para("Suggested pilot: agent-assist copilot on one contact-center program, measured against a "
     "matched control group over 6–8 weeks.")

# ============================================================ 6. HR
h1("6. Use Case 2 — AI-Driven HR Recruitment")
para("Afni hires at high volume, continuously, across three countries — making recruiting an "
     "ideal internal proving ground: Afni controls the data, the risk and the rollout, and the "
     "use case reuses the voice platform for optional pre-screens.")
h2("6.1 Agents across the candidate journey")
bullet("consistent, inclusive job postings.", "JD generation — ")
bullet("résumé parsing and ranking against structured, job-related criteria.", "Sourcing & screening — ")
bullet("chat and optional voice pre-screen (reusing the voice platform).", "Conversational screening — ")
bullet("calendar and ATS automation.", "Scheduling — ")
bullet("structured scoring that assists — never replaces — human interviewers.", "Interview scoring — ")
bullet("continuous adverse-impact monitoring across the funnel.", "Fairness monitor — ")
callout("Non-negotiable principle:",
        "AI assists, humans decide. There is no autonomous rejection of candidates. Every "
        "automated employment decision tool is bias-audited and explainable.")
h2("6.2 Fairness & compliance")
para("The design complies with EEOC guidance, NYC Local Law 144 (bias audits for automated "
     "employment decision tools), the Illinois AI Video Interview Act, the EU AI Act's "
     "high-risk employment provisions, and GDPR. Candidates receive notice and consent, and "
     "decisions remain explainable and human-owned.")
h2("6.3 KPIs")
make_table(["KPI", "Target *"],
           [["Recruiter screening effort", "30–50% reduction"],
            ["Time-to-fill", "Materially shorter"],
            ["Cost-per-hire", "Reduced"],
            ["Candidate experience (NPS)", "Improved"],
            ["Offer-accept rate", "Improved"],
            ["90-day attrition", "Reduced via better matching"]],
           widths=[Inches(3.2), Inches(3.2)])
para("* Illustrative — validated in discovery.", 9, italic=True, color=GRAY)

# ============================================================ 7. LLMOPS LIFECYCLE
h1("7. LLMOps Lifecycle & Toolchain")
para("LLMOps is the operational backbone that turns pilots into dependable production systems. "
     "The lifecycle is a continuous loop: curate data and knowledge → engineer versioned "
     "prompts and agents → evaluate → ship via CI/CD → serve → observe → feed learnings back "
     "into datasets. Governance and Responsible AI wrap the entire loop.")
h2("7.1 Evaluation — the quality gate")
bullet("golden datasets scored by LLM-as-judge plus human review;")
bullet("groundedness / faithfulness scoring for retrieval-augmented answers;")
bullet("regression gates that block promotion when quality drops;")
bullet("red-teaming and safety evaluations before every release;")
bullet("online A/B and shadow testing to validate in production.")
h2("7.2 CI/CD & serving")
para("Prompts, agents and model configurations are versioned in a registry. Releases use "
     "canary or blue-green deployment with instant rollback. Serving runs behind the API "
     "Management gateway with quotas, caching and metering.")
h2("7.3 Toolchain")
make_table(["Capability", "Azure service"],
           [["Prompt & agent authoring", "Prompt flow (Azure AI Foundry)"],
            ["Evaluation", "Azure AI evaluation SDK"],
            ["CI/CD", "GitHub Actions / Azure DevOps"],
            ["Registry", "Model & prompt registry"],
            ["Serving & gateway", "Container Apps / AKS + API Management"],
            ["Tracing", "App Insights + OpenTelemetry (GenAI conventions)"]],
           widths=[Inches(2.6), Inches(3.8)])
para("Unlike traditional MLOps, LLMOps must manage prompts as versioned artifacts, evaluate "
     "non-deterministic outputs, guard against novel failure modes (hallucination, prompt "
     "injection), and treat token cost and latency as first-class release criteria.")

# ============================================================ 8. RESPONSIBLE AI
h1("8. Responsible AI & Governance")
para("Governance follows Microsoft's Responsible AI pillars — fairness, reliability & safety, "
     "privacy & security, inclusiveness, transparency and accountability — operationalized for "
     "Afni through concrete mechanisms rather than principles alone.")
bullet("an AI use-case intake process that risk-tiers every initiative;")
bullet("mandatory human-in-the-loop for consequential decisions;")
bullet("model and system cards for every deployed agent;")
bullet("Content Safety (prompt shields, groundedness, PII, protected material);")
bullet("audit trails, AI incident response, and scheduled red-teaming;")
bullet("an AI governance board with a regular operating cadence.")
h2("8.1 Risk tiering drives controls")
make_table(["Tier", "Examples", "Controls"],
           [["High", "Hiring decisions, collections", "Full HITL, bias audits, legal sign-off"],
            ["Medium", "Customer voice answers", "Guardrails + sampled human QA + monitoring"],
            ["Low", "Internal drafting, summaries", "Standard guardrails + spot checks"]],
           widths=[Inches(1.2), Inches(2.4), Inches(2.8)])

# ============================================================ 9. SECURITY
h1("9. Security, Privacy & Compliance")
bullet("Microsoft Entra ID with least-privilege RBAC and managed identities;", "Identity — ")
bullet("Azure Key Vault; no secrets in code or config;", "Secrets — ")
bullet("VNet with private endpoints; no public egress for sensitive data;", "Network — ")
bullet("at rest and in transit, with optional customer-managed keys;", "Encryption — ")
bullet("honored across US, Mexico and the Philippines;", "Data residency — ")
bullet("Defender for Cloud posture; Purview lineage and DLP; prompt-injection defenses.", "Posture — ")
h2("9.1 Compliance matrix")
make_table(["Framework", "Where it applies", "Primary controls"],
           [["PCI-DSS", "Payment capture", "Pause/mask, tokenization, scope isolation"],
            ["HIPAA", "Healthcare programs", "PHI handling, BAAs, access controls"],
            ["TCPA", "Outbound voice", "Consent, do-not-call, disclosures"],
            ["SOC 2", "Platform", "Control framework, audit evidence"],
            ["GDPR", "PII", "Data-subject rights, minimization"],
            ["EEOC / NYC LL144", "Hiring", "Bias audits, notice, human decision"]],
           widths=[Inches(1.5), Inches(2.0), Inches(2.9)])

# ============================================================ 10. OBSERVABILITY
h1("10. Observability & FinOps")
para("You cannot scale what you cannot see or cost. The platform instruments quality, "
     "groundedness, latency (p50/p95), errors, drift, safety events and token usage, traced "
     "end-to-end with OpenTelemetry GenAI semantic conventions into Application Insights and "
     "Azure Monitor, with dashboards for operations, engineering and governance.")
h2("10.1 FinOps")
bullet("token metering and quotas at the API Management gateway;")
bullet("cost showback per use case and business unit;")
bullet("model right-sizing (GPT-4o vs mini vs open-weight);")
bullet("semantic caching and prompt compression;")
bullet("budget guardrails and anomaly alerts;")
bullet("cost-per-resolved-call and cost-per-screen tracked as first-class KPIs.")
make_table(["Example SLO", "Target"],
           [["Voice turn latency (p95)", "< 1 second"],
            ["Groundedness score", "≥ threshold, gated"],
            ["Availability", "99.9%"],
            ["Safety-event rate", "Below alert threshold"]],
           widths=[Inches(3.2), Inches(3.2)])

# ============================================================ 11. OPERATING MODEL
h1("11. Operating Model & Team")
para("A GenAI Center of Excellence (CoE) operates a federated hub-and-spoke model: the CoE owns "
     "the platform, guardrails and standards, while contact-center and HR 'spokes' own their "
     "use cases and outcomes. Subject-matter experts from Operations, HR, Compliance and "
     "Security embed into delivery pods.")
make_table(["Role", "Focus"],
           [["Executive sponsor", "Funding, priorities, unblock"],
            ["AI product owner", "Backlog, outcomes, adoption"],
            ["GenAI architect (lead)", "Architecture, standards, review"],
            ["Prompt / agent engineers", "Agent design and evaluation"],
            ["LLMOps / MLOps engineers", "CI/CD, serving, observability"],
            ["Data engineers", "Knowledge pipelines, RAG data"],
            ["RAI / governance officer", "Risk tiering, audits, policy"],
            ["Security engineer", "Identity, network, compliance"]],
           widths=[Inches(2.4), Inches(4.0)])

# ============================================================ 12. ROADMAP
h1("12. Implementation Roadmap")
make_table(["Phase", "Timing", "Focus & key deliverables", "Exit criteria"],
           [["Phase 0", "Weeks 0–4", "Landing zone, security baseline, use-case intake, metrics, data access",
             "Foundations live; pilots scoped"],
            ["Phase 1 · Crawl", "Months 1–3", "Agent-assist copilot + HR screening pilots; eval harness; observability baseline",
             "Pilots show measurable value"],
            ["Phase 2 · Walk", "Months 4–7", "Autonomous voice (scoped); HR voice pre-screen + scheduling; online eval; FinOps; CoE stood up",
             "Production use at scale in ≥1 program"],
            ["Phase 3 · Run", "Months 8–12", "Multiple programs/geos; add subrogation & QA analytics; full governance, DR",
             "Enterprise scale; improvement flywheel running"]],
           widths=[Inches(1.3), Inches(1.0), Inches(2.9), Inches(1.6)])
para("Value is delivered from Phase 1 onward; each phase has explicit exit criteria before "
     "the next begins.")

# ============================================================ 13. BUSINESS CASE
h1("13. Business Case & ROI")
para("The value comes from a small number of well-understood levers. All figures below are "
     "illustrative placeholders, to be replaced with Afni's actual volumes, rates and cost "
     "structure during Phase 0 discovery.", italic=True)
make_table(["Value lever", "Use case", "Illustrative impact"],
           [["Call containment / deflection", "Voice AI", "20–40% of eligible calls automated"],
            ["AHT reduction (agent-assist)", "Voice AI", "15–25% shorter handle time"],
            ["QA coverage", "Voice AI", "From ~5% sampled to 100%"],
            ["Recruiter screening effort", "HR", "30–50% fewer manual hours"],
            ["Time-to-fill & cost-per-hire", "HR", "Both materially reduced"],
            ["Attrition (better matching)", "HR", "Lower 90-day attrition"]],
           widths=[Inches(2.4), Inches(1.3), Inches(2.7)])
para("Illustrative payback: 9–15 months, driven primarily by containment and recruiter-time "
     "savings, with QA coverage and attrition as compounding benefits. Under Afni's Gainshare "
     "model, efficiency gains improve shared margin directly.", bold=True, color=NAVY)

# ============================================================ 14. RISKS
h1("14. Risks & Mitigations")
make_table(["Risk", "Category", "Mitigation"],
           [["Hallucination / wrong answers", "Model/quality", "RAG grounding, groundedness gates, HITL, guardrails"],
            ["Compliance breach (TCPA/PCI/EEOC)", "Legal", "Compliance agent, risk-tiering, bias audits, legal sign-off"],
            ["Runaway token cost", "Cost", "APIM metering, right-sizing, caching, budget alerts"],
            ["Quality regression on change", "Delivery", "Golden-set regression gates, canary + rollback"],
            ["Low adoption / change fatigue", "Adoption", "Copilot-first, agent involvement, enablement, clear ROI"],
            ["Vendor / model lock-in", "Strategic", "Gateway abstraction, model catalog, portable orchestration"],
            ["Data leakage / prompt injection", "Security", "Private networking, Content Safety, prompt shields, DLP"]],
           widths=[Inches(2.2), Inches(1.1), Inches(3.1)])

# ============================================================ 15. RECOMMENDATION
h1("15. Recommendation & Next Steps")
para("Approve a four-week Phase 0 to lock the foundations and launch the initial pilots.",
     bold=True, color=NAVY)
bullet("Confirm the two flagship use cases and success metrics with Operations and HR leaders.", "1.  ")
bullet("Stand up the Azure landing zone, security baseline and guardrail policy.", "2.  ")
bullet("Run use-case intake and risk-tiering; secure data access.", "3.  ")
bullet("Launch the agent-assist copilot and HR-screening pilots in Phase 1.", "4.  ")
para("")
para("Evoke Technologies · GenAI Architecture Practice — prepared for Afni, Inc. Confidential.",
     9, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------------------------------------------------------------------------
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "proposal", "Afni-LLMOps-Proposal.docx")
doc.save(out)
print(f"Saved: {out}")



