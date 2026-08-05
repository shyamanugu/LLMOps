#!/usr/bin/env python3
"""Sendable Word document: LLMOps approach for APIX & Hiring Intelligence.
Consolidated approach + activities + as-is/to-be + observability + evaluation + infrastructure.
NO timelines. Requires: python-docx.  Output: document/LLMOps-Approach.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x3A,0x5F); INDIGO=RGBColor(0x2F,0x5C,0x9E); TEAL=RGBColor(0x1E,0x7A,0x72)
GRAY=RGBColor(0x50,0x5A,0x6A); WHITE=RGBColor(0xFF,0xFF,0xFF); MONO="Consolas"
AMBER=RGBColor(0x9A,0x63,0x00); FONT="Calibri"
ROOT=os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc=Document()
n=doc.styles["Normal"]; n.font.name=FONT; n.font.size=Pt(10.5); n.font.color.rgb=RGBColor(0x22,0x28,0x33)
n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.13
for lvl,size,col in [(1,16,NAVY),(2,13,INDIGO),(3,11.5,TEAL)]:
    st=doc.styles[f"Heading {lvl}"]; st.font.name=FONT; st.font.size=Pt(size); st.font.color.rgb=col; st.font.bold=True
    st.paragraph_format.space_before=Pt(12 if lvl==1 else 8); st.paragraph_format.space_after=Pt(4)
def _shade(cell,hexfill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),hexfill); tcPr.append(shd)
def para(t="",size=10.5,bold=False,italic=False,color=None,align=None,sa=6,sb=0):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(sa); p.paragraph_format.space_before=Pt(sb)
    if align: p.alignment=align
    if t:
        r=p.add_run(t); r.font.size=Pt(size); r.bold=bold; r.italic=italic; r.font.name=FONT
        if color: r.font.color.rgb=color
    return p
def bullet(t,lead=None,size=10.5):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
    if lead:
        r=p.add_run(lead); r.bold=True; r.font.name=FONT; r.font.size=Pt(size)
    r2=p.add_run(t); r2.font.name=FONT; r2.font.size=Pt(size)
def h1(t): return doc.add_heading(t,level=1)
def h2(t): return doc.add_heading(t,level=2)
def h3(t): return doc.add_heading(t,level=3)
def code(lines):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.space_before=Pt(2)
    _shade_p(p,"F2F4F7")
    for i,ln in enumerate(lines):
        r=p.add_run(("" if i==0 else "\n")+ln); r.font.name=MONO; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x22,0x2A,0x36)
def _shade_p(p,hexfill):
    pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),hexfill); pPr.append(shd)
def table(headers,rows,widths=None,fill="1F3A5F",fs=9.5):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,hh in enumerate(headers):
        c=t.rows[0].cells[i]; _shade(c,fill); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(2)
        r=p.add_run(hh); r.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(9.5); r.font.name=FONT
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for ci,v in enumerate(row):
            if ri%2==1: _shade(cells[ci],"EEF1F7")
            p=cells[ci].paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(2)
            r=p.add_run(str(v)); r.font.size=Pt(fs); r.font.name=FONT
            if ci==0: r.bold=True; r.font.color.rgb=NAVY
    if widths:
        for ci,w in enumerate(widths):
            for row in t.rows: row.cells[ci].width=w
    para("",sa=4); return t
def callout(title,text):
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; c=t.rows[0].cells[0]; _shade(c,"E9F3F0")
    p=c.paragraphs[0]; r=p.add_run(title+"  "); r.bold=True; r.font.color.rgb=TEAL; r.font.name=FONT; r.font.size=Pt(10.5)
    r2=p.add_run(text); r2.font.name=FONT; r2.font.size=Pt(10.5); para("",sa=4)

print("Building sendable Word document...")
sec=doc.sections[0]; sec.top_margin=Inches(0.9); sec.bottom_margin=Inches(0.8); sec.left_margin=Inches(1.0); sec.right_margin=Inches(1.0)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=fp.add_run("LLMOps Approach — APIX & Hiring Intelligence  ·  Working draft for review"); fr.font.size=Pt(8); fr.font.color.rgb=GRAY; fr.font.name=FONT

# ---- TITLE BLOCK ----
para("WORKING DRAFT FOR REVIEW",11,bold=True,color=TEAL,sa=8)
para("LLMOps: Our Approach",24,bold=True,color=NAVY,sa=2)
para("An operating approach for running Generative AI applications reliably — built around two live use cases "
     "(APIX and Hiring Intelligence) and reusable for any project that follows.",12.5,italic=True,color=GRAY,sa=10)
para("Note on scope: this document focuses on the approach, the activities involved, and how observability and "
     "evaluation will work. Timelines are intentionally left out at this stage. Where current-state details are not "
     "yet confirmed, they are marked as assumptions to validate in discovery.",10,italic=True,color=GRAY,sa=8)
para("LLMOps = Large Language Model Operations. APIX = Afni Performance Intelligence Index.",9.5,italic=True,color=GRAY)

# ---- 1. APPROACH ----
h1("1. How We Approach LLMOps")
para("APIX and Hiring Intelligence already run as agent pipelines built by the product team. Our job is not to "
     "rebuild them. It is to wrap them with an operational layer — the same way a team wraps normal software with "
     "source control, testing, monitoring, and safe releases — and to standardise that layer so any future use case "
     "reuses it instead of starting from scratch.")
callout("The idea in one line:",
        "We add the operational layer (tracking, evaluation, prompt & model control, safe releases, feedback) around "
        "the existing pipelines, make it consistent, and make it reusable.")
para("Both use cases are sequential pipelines — one agent's output feeds the next. They are not agent-to-agent "
     "systems (agents do not freely negotiate with each other). That keeps the operational picture straightforward: "
     "a request flows through a known set of steps, and we can observe and evaluate each step.")
h2("The operational layer, at a glance")
table(["Component","What it gives us"],
      [["Observability","see every step of every request — model calls, tool calls, agent steps"],
       ["Evaluation","measure answer quality with repeatable tests, and block bad changes"],
       ["Prompt management","prompts are versioned, reviewed, and testable — not edited in place"],
       ["Model management","choose and swap models through config, not code changes"],
       ["CI/CD","test and release changes automatically, gradually, and reversibly"],
       ["Guardrails","safety and personal-data checks on inputs and outputs"],
       ["Data / RAG","feed answers from our own documents and data"],
       ["Feedback loop","capture real usage and turn it into improvements"]],
      widths=[Inches(1.8),Inches(4.6)])

# ---- 2. THE TWO USE CASES ----
h1("2. The Two Use Cases as Pipelines")
para("We ground the approach in the two use cases so it is concrete. Pipeline internals below are our current "
     "understanding and should be confirmed with the product team.")
h2("2.1 APIX — Afni Performance Intelligence Index")
para("APIX turns AI-analysed call transcripts into a weekly, per-agent performance report: a composite score out of "
     "100, key performance indicators (KPIs) across sales, retention and customer experience, a four-week trend, "
     "AI-generated coaching recommendations, and risk flags. It runs for multiple programs (Telesales and WCC) with "
     "different scoring criteria.")
para("Pipeline (sequential):", bold=True)
code(["call recording -> speech-to-text transcript (+ metadata: agent, program, outcome)",
      "   -> [1] transcript prep / segmentation",
      "   -> [2] dimension-analysis agents (sales, customer experience, retention, compliance)",
      "   -> [3] extraction agent (escalations, sentiment, sales outcomes)",
      "   -> [4] scoring / aggregation into composite score /100 (program-weighted)",
      "   -> [5] coaching-recommendation agent (practical steps + risk flags)",
      "   -> results stored -> APIX dashboard (managers & coaches)"])
para("Evaluation matters most here for groundedness (coaching must cite what was actually said), scoring agreement "
     "with a human quality reviewer, and consistency/fairness across agents and programs.")
h2("2.2 Hiring Intelligence")
para("Hiring Intelligence supports high-volume recruiting: it parses and ranks résumés against a role, answers "
     "candidate questions, and produces a structured summary and fit score for a human recruiter, who decides.")
para("Pipeline (sequential):", bold=True)
code(["intake / router -> résumé parse & rank (RAG over job description + rubric)",
      "   -> screening Q&A (RAG over role / policy) -> scoring & summary -> human recruiter decides"])
para("It uses tools through MCP (Model Context Protocol) — the applicant tracking system (ATS), the requisition "
     "database, and scheduling. This is where tool-selection evaluation matters: did the agent call the right tool "
     "with the right inputs?")
callout("Why these two together:",
        "APIX exercises groundedness, structured scoring, writing quality and fairness; Hiring exercises tool "
        "selection, retrieval quality and fairness. Between them they cover every evaluation metric group, which is "
        "why the approach generalises to any future use case.")

# ---- 3. AS-IS / TO-BE ----
h1("3. What Exists Today vs What Changes (As-Is / To-Be)")
para("The as-is column is our assumption until we audit the pipelines with the team. Every row is a discovery "
     "question, not a stated fact. We are adding an operational layer around working use cases, not rebuilding them.",
     italic=True)
table(["Area","As-is (to confirm in discovery)","Target (to-be)","Change needed"],
      [["Source control","code in repos; prompts possibly inline","monorepo; prompts/agents/evals versioned","move prompts & agents into Git"],
       ["Prompts","edited in code or a portal, untracked","Git source of truth + registry, reviewed","introduce prompt files + review"],
       ["Models","model names in code","task-aliases in config; swap via change","config-driven model choice"],
       ["Tracing","app logs; no per-step detail","full trace tree (agent/model/tool spans)","add OpenTelemetry tracing"],
       ["Evaluation","manual / spot-check; no gate","golden datasets + automated scoring + CI gate","build the evaluation framework"],
       ["Data / RAG","ad-hoc ingestion","managed ingestion + scheduled refresh","standardise the pipeline"],
       ["Guardrails","minimal","safety + personal-data + human review","add Content Safety + checks"],
       ["Deploy","manual","automated, gated, gradual, reversible","GitHub Actions + gated environments"],
       ["Hosting","to confirm","Azure Container Apps (+ Functions)","containerise the pipeline steps"]],
      widths=[Inches(1.2),Inches(2.2),Inches(2.0),Inches(1.6)],fs=8.5)

# ---- 4. ACTIVITIES ----
h1("4. The Activities Involved")
para("The work is organised into workstreams. This is the order of work and its dependencies — not a schedule. No "
     "dates are proposed at this stage.")
def ws(letter,name,goal,acts):
    h3(f"{letter}. {name}")
    para(goal, italic=True, color=GRAY, sa=3)
    for a in acts: bullet(a)
ws("A","Discovery & current-state assessment","Understand what exists before changing anything.",
   ["Inventory the APIX and Hiring pipelines: agents, prompts, models, tools, data sources.",
    "Review current logging and any evaluation in place.",
    "Confirm the as-is table and identify the real gaps."])
ws("B","Foundation","Stand up the shared base the pipelines will run on.",
   ["GitHub repository and folder structure (/prompts, /agents, /evals, /src, /pipelines, /infra, /dashboards).",
    "Azure landing zone: identity (Entra ID), secrets (Key Vault), gateway (API Management), model deployments."])
ws("C","Instrumentation & observability","Make every step of every request visible.",
   ["Add tracing (OpenTelemetry) to both pipelines: agent, model-call and tool-call spans.",
    "Stand up Azure Application Insights (system of record) and self-hosted Langfuse (LLM-specific view).",
    "Define exactly what is captured per request, model call, tool call, and session (see section 5)."])
ws("D","Evaluation framework  (priority — runs early and continuously)","Measure quality and stop bad changes.",
   ["Build golden datasets per use case and per program (Telesales and WCC differ).",
    "Implement evaluators: Ragas and DeepEval for retrieval and writing quality; custom Python for tool selection and scoring.",
    "Wire an automated quality gate into the release pipeline; add online sampling and human review (see section 6)."])
ws("E","Prompt & model management","Treat prompts and model choices as reviewed, testable configuration.",
   ["Move prompts to Git with a runtime registry (labels for production and staging).",
    "Introduce model task-aliases so no model name is hard-coded; any change goes through the quality gate."])
ws("F","CI/CD & release","Ship changes safely and reversibly.",
   ["GitHub Actions pipelines with federated login to Azure (no stored keys).",
    "Gated environments (dev, test, production); gradual (canary) release; automatic rollback."])
ws("G","Data & knowledge pipelines","Feed answers from our own data, kept fresh.",
   ["Hiring retrieval ingestion (job descriptions, rubrics, policy).",
    "APIX transcript and metadata flow; scheduled and change-driven refresh."])
ws("H","Guardrails & governance","Keep it safe, compliant, and fair.",
   ["Content Safety and personal-data (PII) handling; fairness checks for hiring ranking and APIX consistency.",
    "Human-in-the-loop for consequential outputs."])
ws("I","Feedback & improvement loop","Turn real usage into steady improvement.",
   ["Capture coach and recruiter feedback; analytics dashboards.",
    "Triage weak results, add them to the golden datasets, fix, re-evaluate, and ship."])
para("Sequencing: A–C are foundational; D (evaluation) starts early and runs throughout; E–I layer in as we go.", bold=True, color=NAVY)

# ---- 5. OBSERVABILITY ----
h1("5. How Observability Works")
para("A request is a trace. Everything that happens inside it is a nested span (a step with its own recorded "
     "detail). For a pipeline, the trace looks like this:")
code(["Request (trace)  ── one call analysed / one candidate screened",
      "  └─ Agent 1 (span)  ── e.g. dimension analysis / résumé rank",
      "       ├─ Model call (span)",
      "       └─ Tool call (span)",
      "  └─ Agent 2 (span)  ── e.g. scoring / screening",
      "  └─ Final output"])
para("Each span records its own inputs, outputs, timing and cost, and rolls up to its parent — so the total time and "
     "cost of the whole request add up automatically. One trace id ties it all together, so a bad answer can be "
     "followed back to the exact step.")
h2("What we capture at each level")
table(["Level","What we record"],
      [["Request (trace)","use case, program, input id (call/candidate), final output ref, status, total latency, total tokens, total cost, whether a human intervened"],
       ["Agent step (span)","agent name and version, its input and output, which model(s) and tool(s) it used, step latency, tokens, cost, and which step it handed off to"],
       ["Model call (span)","model name and version, prompt id and version, the prompt and the answer (hashed if it contains personal data), tokens in/out, cost, latency, temperature, finish reason, cache hit"],
       ["Tool call (span)","tool name, which MCP server, the inputs, the result, success or error, latency — and the expected tool and whether the correct tool was chosen (this feeds tool-selection evaluation)"],
       ["Session","for multi-turn chats (e.g. Hiring screening): links the turns, conversation id, user id (hashed), number of turns, outcome"],
       ["Feedback","thumbs up/down with a reason, coach edits to a report, recruiter overrides — linked back to the trace id"]],
      widths=[Inches(1.5),Inches(4.9)],fs=9)
para("This directly answers the four questions raised: what is tracked for every request (the trace row), how model "
     "calls are tracked (the model-call span), how tool calls are tracked (the tool-call span, including the "
     "correct-tool flag), and how agent sessions are monitored (the session id linking the agent spans).")
h2("How it is built on Azure")
bullet("instrumented with OpenTelemetry (an open standard) using its Generative AI conventions.", "Standard — ")
bullet("Azure Application Insights + Log Analytics as the system of record (data stays in our tenant).", "Record of truth — ")
bullet("self-hosted Langfuse (on Azure Container Apps + Azure Database for PostgreSQL) for the LLM-specific view — cost per model, prompt versions, per-trace scores.", "LLM lens — ")
bullet("Azure AI Foundry tracing links evaluation scores back to the exact trace.", "Eval link — ")
para("Personal data (candidate or customer information) in prompts is hashed or redacted before it is stored in traces.", italic=True, color=GRAY)

# ---- 6. EVALUATION ----
h1("6. How Evaluation Works")
para("This is where we will spend the most attention. LLM output is not deterministic — the same input can produce "
     "different wording, and a fluent answer can still be wrong. So we do not check output == expected; we run a set "
     "of test cases on every change and score the output on several dimensions.")
h2("6.1 Metric groups")
para("We group metrics by what is being judged, so use cases beyond retrieval are still covered.")
table(["Group","Example metrics","Judges…"],
      [["Retrieval / RAG quality","groundedness, context relevance, retrieval precision/recall","is it backed by our data?"],
       ["Generation / writing quality","coherence, tone, completeness","how well it reads"],
       ["Task execution / agentic","tool-selection accuracy, correct action, steps taken","whether it did the right thing"],
       ["Safety / compliance / fairness","unsafe-content rate, PII leakage, bias","is it safe and fair?"],
       ["Operational","latency, cost, tokens per request","is it fast and affordable?"]],
      widths=[Inches(1.9),Inches(2.9),Inches(1.6)])
para("Why writing quality and task execution are separate groups: writing quality judges how the answer reads; task "
     "execution judges whether the system did the right thing — for example, whether it called the correct tool. "
     "These are independent: an answer can read well and still take the wrong action. Some metrics overlap (coherence "
     "could sit under retrieval or writing); we assign each metric to one group so we do not double-count.")
h2("6.2 Tool-selection evaluation")
para("When an MCP tool server exposes several tools, choosing the wrong tool but still returning an answer is "
     "unreliable. We test this directly. For each test case we know the correct tool (and inputs). We run the agent, "
     "read the tool it actually chose from the trace, and compare.")
para("We measure: tool-selection accuracy; per-tool precision and recall; wrong-tool rate; calling a tool when none "
     "was needed; missing a tool it should have used; and argument correctness (right tool, but were the inputs "
     "right?). Retrieval frameworks like Ragas and DeepEval do not cover this, so it is a small custom Python check:")
code(["for case in golden_tool_cases:            # each has input + expected_tool (+ expected_args)",
      "    trace = run_agent(case.input)",
      "    chosen = trace.tool_calls[0].name      # what the agent actually called",
      "    args_ok = compare_args(trace.tool_calls[0].args, case.expected_args)",
      "    record(correct = (chosen == case.expected_tool), args_ok = args_ok)",
      "# report accuracy, per-tool precision/recall, wrong-tool / missing-tool / unnecessary-call rates"])
h2("6.3 The tools we use")
para("No single tool does everything. We recommend a mix.")
table(["Tool","Covers","Open source?","Use it for"],
      [["Ragas","RAG metrics (groundedness, relevance)","Yes","APIX groundedness, Hiring retrieval"],
       ["DeepEval","broad LLM eval, custom metrics, CI-friendly","Yes","the quality gate, writing quality"],
       ["Custom Python","tool selection, scoring vs labels, extraction","Yes (our code)","agent and tool behaviour"],
       ["LLM-as-judge (rubric)","subjective quality","depends","coaching usefulness, summaries"],
       ["Azure AI Foundry evaluations","built-in + custom, links to traces","No (Azure)","staying inside Azure"],
       ["promptfoo","config-driven checks, red-team","Yes","fast checks in CI"],
       ["LangSmith","evaluation + observability platform","No — licensed","only if we standardise on it (cost)"]],
      widths=[Inches(1.7),Inches(2.5),Inches(1.1),Inches(1.9)],fs=8.5)
h2("6.4 How and when it runs")
bullet("golden datasets per use case and per program; runs on every change and blocks a release if quality drops past a threshold.", "Offline (quality gate) — ")
bullet("sample a share of live traffic, score it in the background, and alert if quality drifts.", "Online (production) — ")
bullet("coaches and recruiters give feedback; experts review a sample; findings become new test cases.", "Human review — ")
para("We evaluate each pipeline step and the final output. A pipeline can look fine end-to-end while one step quietly "
     "gets worse — checking both catches that.", bold=True, color=NAVY)
h2("6.5 Golden datasets")
para("A golden dataset is a saved, versioned set of test cases: an input (plus any context), the expected answer or "
     "a grading rule, and metadata (intent, program, source). They come from three places — experts write realistic "
     "cases, we mine anonymised real traffic from the traces, and we generate synthetic cases that a person reviews. "
     "Start with roughly 50–200 per use case and program, and grow the set from production feedback. Example record:")
code(["{",
      "  \"id\": \"apix-telesales-014\",",
      "  \"input\": {\"transcript_id\": \"c-88421\", \"program\": \"telesales\"},",
      "  \"grading_criteria\": {",
      "    \"must_cite_evidence\": true,",
      "    \"expected_score_band\": [70, 85],",
      "    \"must_flag\": [\"missed_upsell\"]",
      "  },",
      "  \"metadata\": {\"program\": \"telesales\", \"source\": \"sme_authored\"}",
      "}"])

# ---- 7. INFRASTRUCTURE ----
h1("7. Infrastructure & Azure Hosting (Proposed)")
para("This is the target setup. Sequencing is covered in section 4; no dates are proposed here.")
h2("Hosting the pipelines")
table(["Option","Fit","Note"],
      [["Azure Container Apps (recommended)","each pipeline step as a service; scales to zero","best general fit for pipeline services"],
       ["Azure Functions","event-driven triggers (new transcript, new candidate)","good for APIX batch/event jobs"],
       ["Foundry Agent Service","managed hosted agents; less to run","consider as it matures"]],
      widths=[Inches(2.4),Inches(2.6),Inches(1.4)])
para("Recommendation: Azure Container Apps for the pipeline services, with Azure Functions for event triggers. "
     "Foundry Agent Service can be adopted later for hosted agents.")
h2("The shared platform (bill of services)")
table(["Layer","Azure service","Purpose"],
      [["Models / AI","Azure OpenAI, Content Safety","the models; safety checks"],
       ["Knowledge / RAG","Azure AI Search","retrieval for Hiring; optional transcript search"],
       ["Data / state","Cosmos DB / Azure SQL, Blob Storage","agent state, APIX scores, transcripts, datasets"],
       ["Analytics (later)","Microsoft Fabric / OneLake","telemetry lake, dashboards, training-data curation"],
       ["Gateway / compute","API Management, Container Apps, Functions","one entry point; run the services"],
       ["Observability","Azure Monitor, Application Insights, self-hosted Langfuse","tracing, dashboards, LLM-specific view"],
       ["Web app (APIX)","App Service / Static Web Apps","the dashboard front end"],
       ["Security / identity","Entra ID, Key Vault, Private Endpoints, Purview, Defender","identity, secrets, network, governance"],
       ["CI/CD","GitHub + GitHub Actions (federated login)","test and release automatically"]],
      widths=[Inches(1.5),Inches(2.6),Inches(2.3)],fs=8.8)
para("Environments: separate dev, test and production on a landing zone. Shared platform vs per-use-case: everything "
     "above is built once and shared; per use case we add only the prompts, agents, tools, and golden datasets.")

# ---- 8. WHAT WE NEED TO PROCEED ----
h1("8. What We Need to Proceed")
para("No dates yet — this is what unblocks the work when we are ready:")
bullet("access to review the APIX and Hiring pipelines with the product team (to confirm the as-is).")
bullet("Azure subscription access and model quota.")
bullet("access to the relevant data sources (transcripts, job descriptions, rubrics).")
bullet("subject-matter expert time to help build the first golden datasets.")
para("")
para("Prepared for internal review. Feedback and refinement welcome — this is a starting point for the approach, "
     "not a final plan.",9.5,italic=True,color=GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)

out=os.path.join(ROOT,"document","LLMOps-Approach.docx"); doc.save(out); print(f"Saved: {out}")
